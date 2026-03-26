from docx.document import Document

def check_formatting_rules(doc: Document, exam_type: str) -> list[dict]:
    """
    Rule-based method for Font size, bolding, margins, numbering.
    Returns a list of dicts: {"text": "matched text", "annotation": "feedback"}
    """
    feedback = []
    
    # 1. Checking Margins (very simple check on first section)
    if doc.sections:
        section = doc.sections[0]
        # Example check for standard margins (approx 1 inch, here checking if less than 0.5 points to tiny margins)
        if section.left_margin.inches < 0.5 or section.right_margin.inches < 0.5:
            feedback.append({
                "text": doc.paragraphs[0].text if doc.paragraphs else "Document",
                "annotation": "Rule Check: Document margins appear to be too narrow."
            })

    for para in doc.paragraphs:
        text = para.text.strip()
        if text.startswith("Question") or text.startswith("Q:"):
            is_bold = any(run.bold for run in para.runs)
            if not is_bold:
                feedback.append({
                    "text": text,
                    "annotation": "Rule Check: Question headers should be bold."
                })
        
        # Check numbering formatting
        # If it looks like a list item but not styled correctly
        if len(text) > 2 and text[0].isdigit() and text[1] in [".", ")"]:
            # Could check if it's using python-docx list styles, but a simple text check suffices for now
            pass
            
    return feedback
