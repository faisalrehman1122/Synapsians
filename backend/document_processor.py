import io
import markdownify
from docx import Document
from rules_engine import check_formatting_rules
import progress
import re


def extract_questions_and_map_paragraphs(doc):
    """
    Builds LLM question chunks AND the exact paragraph mapping in a single pass.
    This guarantees that the question ID the LLM evaluates is physically bound
    to the same paragraphs used for comment insertion — no index shift possible.
    """
    exam_questions = []
    paragraphs_by_qidx = {}

    current_qidx = 0
    current_chunk_text = ""
    current_paras = []

    for para in doc.paragraphs:
        if "--" in para.text:
            # Save the current chunk if it contains a real question
            if current_chunk_text.strip() and ("Typ:" in current_chunk_text or "Modus:" in current_chunk_text):
                match = re.search(r'(Typ|Modus):\s*(PickS|Kprim|TypA)', current_chunk_text, re.IGNORECASE)
                q_type = match.group(2) if match else "TypA"

                exam_questions.append({
                    "id": current_qidx,
                    "type": q_type,
                    "markdown": current_chunk_text.strip()
                })
                paragraphs_by_qidx[current_qidx] = current_paras
                current_qidx += 1

            # Reset for the next question
            current_chunk_text = ""
            current_paras = []
        else:
            # Extract text with bold markdown tags
            para_md = ""
            for run in para.runs:
                text = run.text
                if run.bold and text.strip():
                    para_md += f"**{text}**"
                else:
                    para_md += text

            current_chunk_text += para_md + "\n"
            current_paras.append(para)

    # Save the last chunk if the document doesn't end with "--"
    if current_chunk_text.strip() and ("Typ:" in current_chunk_text or "Modus:" in current_chunk_text):
        match = re.search(r'(Typ|Modus):\s*(PickS|Kprim|TypA)', current_chunk_text, re.IGNORECASE)
        q_type = match.group(2) if match else "TypA"
        exam_questions.append({
            "id": current_qidx,
            "type": q_type,
            "markdown": current_chunk_text.strip()
        })
        paragraphs_by_qidx[current_qidx] = current_paras

    return exam_questions, paragraphs_by_qidx


async def process_exam_document(file_bytes: bytes) -> bytes:
    doc = Document(io.BytesIO(file_bytes))
    
    # 1. Read exam type from header
    exam_type = "Unknown"
    for section in doc.sections:
        for header_para in section.header.paragraphs:
            if header_para.text.strip():
                exam_type = header_para.text.strip()
                break
        if exam_type != "Unknown":
            break
            
    # Fallback to first paragraph if no header
    if exam_type == "Unknown" and doc.paragraphs:
        exam_type = doc.paragraphs[0].text.strip()
    
    # 2. Extract questions AND map paragraphs in a single pass (prevents index-shift bugs)
    progress.current_status["message"] = "Splitting document into markdown questions..."
    progress.current_status["progress"] = 15
    progress.current_status["phase"]    = "parsing"
    
    import asyncio
    
    exam_questions, paragraphs_by_qidx = extract_questions_and_map_paragraphs(doc)
    
    # 3. Rule-based format check
    progress.current_status["message"] = "Running Python Rule-Based Formatting Engine..."
    progress.current_status["progress"] = 25
    progress.current_status["phase"]    = "parsing"
    rule_feedback = check_formatting_rules(doc, exam_type)
    
    # 4. LLM Approach for questions using parallel execution
    from llm_engine import process_exam_in_parallel
    
    # Run the synchronous parallel processor in a separate thread so we don't block FastAPI
    total_q = len(exam_questions)
    progress.current_status["message"] = f"Sending {total_q} questions to AI..."
    progress.current_status["progress"] = 30
    progress.current_status["phase"]    = "processing"
    
    llm_results = await asyncio.to_thread(process_exam_in_parallel, exam_questions)
    
    # Extract feedback format to match downstream Word Comment insertion logic
    progress.current_status["message"] = "Collating rules and LLM annotations..."
    progress.current_status["progress"] = 85
    progress.current_status["phase"]    = "collating"
    llm_feedback = []
    for res in llm_results:
        if res.get("success") and "feedback" in res:
            q_idx = res.get("id", -1)
            for item in res["feedback"].get("feedback_comments", []):
                llm_feedback.append({
                    "q_idx": q_idx,
                    "text": item.get("exact_quote", ""),
                    "annotation": item.get("comment", "")
                })
            
    # Normalize rules formatting feedback
    for rf in rule_feedback:
        rf["q_idx"] = -1
        
    # Combine feedback
    all_feedback = rule_feedback + llm_feedback
    
    # 5. Apply comments using native doc.add_comment()
    progress.current_status["message"] = f"Writing {len(all_feedback)} comments into document..."
    progress.current_status["progress"] = 90
    progress.current_status["phase"]    = "collating"

    seen_paras = set()
    
    for feedback_item in all_feedback:
        target_text = feedback_item["text"].strip()
        annotation = feedback_item["annotation"].strip()
        q_idx = feedback_item.get("q_idx", -1)
        added = False
        
        # Ensure we're finding non-empty text (and skip single characters to avoid reckless matching)
        if not target_text or len(target_text) < 2:
            continue
            
        search_paras = doc.paragraphs
        if q_idx != -1 and q_idx in paragraphs_by_qidx:
            search_paras = paragraphs_by_qidx[q_idx]
            
        for para in search_paras:
            if target_text in para.text:
                if para.runs:
                    try:
                        # Anchor to the first run to prevent XML corruption across complex runs
                        doc.add_comment(para.runs[0], text=annotation, author="Exam Evaluator")
                    except Exception as e:
                        print(f"Failed to add comment to paragraph: {e}")
                added = True
                break
                
        # If not found in main paragraphs, try tables
        if not added:
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if target_text in para.text:
                                if para.runs:
                                    try:
                                        # Anchor to the first run
                                        doc.add_comment(para.runs[0], text=annotation, author="Exam Evaluator")
                                    except Exception as e:
                                        print(f"Failed to add comment to table paragraph: {e}")
                                added = True
                                break
                        if added: break
                    if added: break
                if added: break
                
    # Save modified document
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.read()

# Make sure to import docx for the fallback RGBColor
import docx
