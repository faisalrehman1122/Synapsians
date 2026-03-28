import os
import json
from docx import Document
from dotenv import load_dotenv

# Load env variables so the LLM engine can authenticate with Azure
load_dotenv("../.env")
import llm_engine

DATA_DIR = "../data/Selektion Klausuren Hackathon ki"
OUT_DIR = "../data/ft_data"

os.makedirs(OUT_DIR, exist_ok=True)

# Dynamically get docx files
all_files = [f for f in os.listdir(DATA_DIR) if f.endswith(".docx") and not f.startswith("~$")]

# Split into train/val
training_files = all_files[:4]
validation_files = all_files[4:]

def extract_questions_from_docx(filepath):
    doc = Document(filepath)
    exam_type = "Unknown"
    for section in doc.sections:
        for header_para in section.header.paragraphs:
            if header_para.text.strip():
                exam_type = header_para.text.strip()
                break
        if exam_type != "Unknown":
            break
            
    if exam_type == "Unknown" and doc.paragraphs:
        exam_type = doc.paragraphs[0].text.strip()

    full_text = "\n".join([p.text for p in doc.paragraphs])
    questions_raw = full_text.split("--")
    
    exam_questions = []
    for idx, q in enumerate(questions_raw):
        if q.strip():
            exam_questions.append({
                "id": idx,
                "type": exam_type,
                "markdown": q.strip()
            })
    return exam_questions

def create_jsonl(files, output_file):
    if not files:
        return
        
    all_questions = []
    for f in files:
        filepath = os.path.join(DATA_DIR, f)
        print(f"Extracting {f}...")
        all_questions.extend(extract_questions_from_docx(filepath))

    print(f"Processing {len(all_questions)} total questions for {output_file} via Azure OpenAI...")
    
    # Process them in parallel using the existing engine
    results = llm_engine.process_exam_in_parallel(all_questions)
    
    out_path = os.path.join(OUT_DIR, output_file)
    with open(out_path, "w", encoding="utf-8") as out:
        for question_data, result in zip(all_questions, results):
            if not result.get("success"):
                print(f"Failed to get LLM response for question {question_data['id']}")
                continue
                
            system_prompt = llm_engine.build_system_prompt(question_data["type"])
            user_prompt = question_data["markdown"]
            assistant_response = json.dumps(result.get("feedback", {}), ensure_ascii=False)
            
            jsonl_line = {
                "messages": [
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                    {"role": "assistant", "content": assistant_response}
                ]
            }
            out.write(json.dumps(jsonl_line, ensure_ascii=False) + "\n")
    print(f"Saved to {out_path}")

if __name__ == "__main__":
    print("Generating Training Data...")
    create_jsonl(training_files, "training.jsonl")

    print("\nGenerating Validation Data...")
    create_jsonl(validation_files, "validation.jsonl")

    print("\nDone! Fine-tuning pairs saved to data/ft_data/")
